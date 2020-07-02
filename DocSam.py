# -*- coding: utf-8 -*-
"""
Created on Sun Sep 22 13:12:10 2019

@author: ArunMukesh
"""
'''
go to anaconda cmp open it as administrator
pip install python-docx

'''
from docx import Document
import cx_Oracle
def readFile():
    try:
        lines =[]
        path = input("Please enter the docx's path : ")
        #document = Document(r'C:\Users\ArunM\Desktop\Sam.docx')
        document = Document(path)
        for line in document.paragraphs:
            print(line.text)
            lines.append(line.text)
        return lines
    except Exception as e:
        print(e.__str__)
            

try :
    '''
    Oracle db connect 
    '''
    ips = input("Please enter the ip : ")
    userName = input("Please enter the user name : ")
    pwd = input("Please enter the password : ")
    #conn = cx_Oracle.connect('hr/hr@localhost')
    location = userName+'/'+pwd+'@'+ips
    #print(location)
    conn = cx_Oracle.connect(location)
    cursor = conn.cursor()
    print(" connected to db sucessfully " )
    '''
    sql insert part
    
    create sequence
    
    create sequence FileToTable_seq.nextval start with 1 increment by 1 minvalue 1 maxvalue 99;
    '''
    lines = readFile()
    if not lines:
        raise Exception ("Unable to find the file")
    for line in lines:
        a = line.split(' ')
        name = a[0]
        age = int(a[1])
                
        if a:
            try :
                print ("Name : {} Age :{}".format(name,age) )
                #b = cursor.execute("insert into FileToTable(KEY_1,name,age) values (FileToTable_seq.nextval,?,? )",(name,age))
                cursor.execute("insert into FileToTable(KEY_1,name,age) values(FileToTable_seq.nextval,:name,:age)",{'name':name , 'age':age})
                conn.commit()
                print("Updated successfully...!")
            except cx_Oracle.DatabaseError as bs :
                print ("exception error ",bs)
                conn.rollback()    
            except :
                conn.rollback()
        else :
            print("Issue in reading file  or in data mapping")
    
    
    '''
    Closing existing connection 
    '''
    conn.close()

except cx_Oracle.DatabaseError as e:
    print(" this is an exception ",e)
    conn.close()
except :
    conn.close()
  